import os
from docx import Document
import pdfplumber


def _parse_pdf(file_path):
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[第{index}页]\n{text.strip()}")
    return "\n\n".join(pages).strip()


def _parse_docx(file_path):
    document = Document(file_path)
    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _parse_text(file_path):
    encodings = ["utf-8", "utf-8-sig", "gb18030", "latin-1"]
    last_error = None
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as file:
                return file.read().strip()
        except UnicodeDecodeError as error:
            last_error = error
    raise UnicodeDecodeError("unknown", b"", 0, 1, f"无法解码文本文件：{last_error}")


def parse_resume_file(file_path):
    if not file_path:
        return {"success": False, "error": "file_path不能为空", "text": ""}
    expanded_path = os.path.abspath(os.path.expanduser(file_path))
    if not os.path.exists(expanded_path):
        return {"success": False, "error": f"文件不存在：{expanded_path}", "text": "", "file_path": expanded_path}
    if not os.path.isfile(expanded_path):
        return {"success": False, "error": f"路径不是文件：{expanded_path}", "text": "", "file_path": expanded_path}
    ext = os.path.splitext(expanded_path)[1].lower()
    try:
        if ext == ".pdf":
            text = _parse_pdf(expanded_path)
            file_type = "pdf"
        elif ext == ".docx":
            text = _parse_docx(expanded_path)
            file_type = "docx"
        elif ext in [".txt", ".md", ".text"]:
            text = _parse_text(expanded_path)
            file_type = "text"
        else:
            return {"success": False, "error": f"暂不支持的文件类型：{ext}", "text": "", "file_path": expanded_path}
        return {
            "success": True,
            "file_path": expanded_path,
            "file_type": file_type,
            "char_count": len(text),
            "text": text,
            "warnings": [] if text.strip() else ["文件已读取，但未提取到有效文本"],
        }
    except Exception as error:
        return {"success": False, "error": str(error), "text": "", "file_path": expanded_path}
